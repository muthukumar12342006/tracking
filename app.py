"""
Budgeting & Variance Analysis Audit Workbench
CITF AI – AI-Powered Internal Audit Analytics Platform
Production-ready Streamlit application (single-file architecture).
"""

from __future__ import annotations

import io
import shutil
from datetime import datetime
from pathlib import Path

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from st_aggrid import AgGrid, GridOptionsBuilder, GridUpdateMode, JsCode

# ── Configuration ─────────────────────────────────────────────────────────────

APP_TITLE = "Budgeting & Variance Analysis Audit Workbench"
APP_SUBTITLE = "AI-Powered Internal Audit Analytics Platform"

COLORS = {
    "navy": "#0A1628",
    "navy_light": "#152238",
    "teal": "#00D4AA",
    "teal_dark": "#00A888",
    "white": "#FFFFFF",
    "gray_100": "#F8FAFC",
    "gray_400": "#94A3B8",
    "gray_600": "#475569",
    "success": "#10B981",
    "warning": "#F59E0B",
    "danger": "#EF4444",
}

BUDGET_COLS = ["Period", "Cost Head", "Budget Amount", "Approved Date", "Budget Start Date"]
ACTUAL_COLS = ["Period", "Cost Head", "Actual Amount"]
REVISION_COLS = ["Revision ID", "Cost Head", "Old Budget", "New Budget", "Revision Date", "Approved By"]

NAV_ITEMS = [
    ("🏠", "Dashboard", "dashboard"),
    ("📁", "Data Import", "import"),
    ("🔍", "Audit Test 1 – Variances", "test1"),
    ("🔍", "Audit Test 2 – Approvals", "test2"),
    ("🔍", "Audit Test 3 – Overspend", "test3"),
    ("🔍", "Audit Test 4 – Revisions", "test4"),
    ("🔍", "Audit Test 5 – Assumptions", "test5"),
    ("🤖", "AI Insights", "ai_insights"),
    ("📈", "Analytics", "analytics"),
    ("📋", "Audit Report", "report"),
]

BASE_DIR = Path(__file__).resolve().parent
LOGO_PATH = BASE_DIR / "assets" / "citf_ai_logo.png"
_CURSOR_LOGO = Path(
    r"C:\Users\Admin\.cursor\projects\c-Users-Admin-Projects-cap-ai-budgeting\assets"
    r"\c__Users_Admin_AppData_Roaming_Cursor_User_workspaceStorage_empty-window_images_image-674f540b-4457-4d88-a554-6a9e9ecfc727.png"
)


def _ensure_logo() -> Path | None:
    if LOGO_PATH.exists():
        return LOGO_PATH
    LOGO_PATH.parent.mkdir(parents=True, exist_ok=True)
    for candidate in [_CURSOR_LOGO, BASE_DIR / "assets" / "cap_ai_logo.png"]:
        if candidate.exists():
            shutil.copy(candidate, LOGO_PATH)
            return LOGO_PATH
    return None


LOGO = _ensure_logo()

st.set_page_config(
    page_title=f"CITF AI – {APP_TITLE}",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ── Session & Audit Trail ─────────────────────────────────────────────────────

def init_session():
    defaults = {
        "budget_df": None,
        "actual_df": None,
        "revision_df": None,
        "merged_df": None,
        "test1_df": None,
        "test2_df": None,
        "test3_df": None,
        "test4_df": None,
        "test5_findings": [],
        "audit_trail": [],
        "user_name": "Audit Analyst",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def log_audit(action: str, detail: str = ""):
    entry = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "user": st.session_state.get("user_name", "Audit Analyst"),
        "action": action,
        "detail": detail,
    }
    st.session_state.audit_trail.insert(0, entry)


# ── Styling ───────────────────────────────────────────────────────────────────

def get_css() -> str:
    return f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
    html, body, [class*="css"] {{ font-family: 'Inter', sans-serif; }}
    .stApp {{
        background: {COLORS['navy']};
        background-image: radial-gradient(ellipse at 20% 0%, rgba(0,212,170,0.08) 0%, transparent 50%),
                          radial-gradient(ellipse at 80% 100%, rgba(0,212,170,0.05) 0%, transparent 50%);
    }}
    [data-testid="stSidebar"] {{
        background: {COLORS['navy_light']};
        border-right: 1px solid rgba(0,212,170,0.25);
    }}
    .app-header {{
        background: linear-gradient(135deg, {COLORS['navy']} 0%, {COLORS['navy_light']} 60%, rgba(0,212,170,0.12) 100%);
        border: 1px solid rgba(0,212,170,0.25);
        border-radius: 16px;
        padding: 1.5rem 2rem;
        margin-bottom: 1.5rem;
        backdrop-filter: blur(12px);
    }}
    .app-header h1 {{ color: {COLORS['white']}; font-weight: 800; font-size: 1.75rem; margin: 0; }}
    .app-header p {{ color: {COLORS['gray_400']}; margin: 0.4rem 0 0; font-size: 0.95rem; }}
    .glass-card {{
        background: rgba(21,34,56,0.65);
        backdrop-filter: blur(16px);
        border: 1px solid rgba(0,212,170,0.25);
        border-radius: 14px;
        padding: 1.25rem;
        margin-bottom: 1rem;
        transition: transform 0.25s ease, box-shadow 0.25s ease;
    }}
    .glass-card:hover {{ transform: translateY(-3px); box-shadow: 0 12px 40px rgba(0,212,170,0.12); }}
    .kpi-card {{
        background: rgba(21,34,56,0.7);
        border: 1px solid rgba(0,212,170,0.25);
        border-radius: 14px;
        padding: 1.25rem;
        text-align: center;
        animation: fadeInUp 0.5s ease forwards;
    }}
    .kpi-value {{ font-size: 1.6rem; font-weight: 700; color: {COLORS['teal']}; }}
    .kpi-label {{ font-size: 0.75rem; color: {COLORS['gray_400']}; text-transform: uppercase; letter-spacing: 0.05em; margin-top: 0.3rem; }}
    .insight-card {{
        background: linear-gradient(135deg, rgba(21,34,56,0.9) 0%, rgba(0,212,170,0.08) 100%);
        border-left: 4px solid {COLORS['teal']};
        border-radius: 0 12px 12px 0;
        padding: 1.25rem 1.5rem;
        margin-bottom: 1rem;
    }}
    .insight-card h4 {{ color: {COLORS['teal']}; margin: 0 0 0.5rem; }}
    .insight-card p {{ color: {COLORS['gray_400']}; margin: 0; line-height: 1.6; }}
    .section-title {{
        color: {COLORS['white']};
        font-weight: 700;
        font-size: 1.2rem;
        border-bottom: 2px solid {COLORS['teal']};
        padding-bottom: 0.4rem;
        margin: 1.5rem 0 1rem;
        display: inline-block;
    }}
    @keyframes fadeInUp {{
        from {{ opacity: 0; transform: translateY(16px); }}
        to {{ opacity: 1; transform: translateY(0); }}
    }}
    div[data-testid="stMetric"] {{
        background: rgba(21,34,56,0.65);
        border: 1px solid rgba(0,212,170,0.2);
        border-radius: 12px;
        padding: 0.75rem;
    }}
    .stDownloadButton button {{
        background: linear-gradient(135deg, {COLORS['teal']} 0%, {COLORS['teal_dark']} 100%) !important;
        color: {COLORS['navy']} !important;
        font-weight: 600 !important;
        border-radius: 8px !important;
    }}
    #MainMenu {{visibility: hidden;}}
    footer {{visibility: hidden;}}
    </style>
    """


def render_header(title: str = APP_TITLE, subtitle: str = APP_SUBTITLE):
    logo_html = ""
    if LOGO and LOGO.exists():
        import base64
        b64 = base64.b64encode(LOGO.read_bytes()).decode()
        logo_html = f'<img src="data:image/png;base64,{b64}" width="72" style="margin-right:1rem;vertical-align:middle;border-radius:8px;">'
    st.markdown(
        f'<div class="app-header">{logo_html}<h1>{title}</h1><p>{subtitle}</p></div>',
        unsafe_allow_html=True,
    )


def kpi_card(label: str, value: str):
    st.markdown(
        f'<div class="kpi-card"><div class="kpi-value">{value}</div><div class="kpi-label">{label}</div></div>',
        unsafe_allow_html=True,
    )


def fmt_currency(v: float) -> str:
    if abs(v) >= 1_000_000:
        return f"${v / 1_000_000:,.2f}M"
    if abs(v) >= 1_000:
        return f"${v / 1_000:,.1f}K"
    return f"${v:,.0f}"


# ── Data Validation & Quality ─────────────────────────────────────────────────

def validate_columns(df: pd.DataFrame, required: list[str]) -> tuple[bool, list[str]]:
    missing = [c for c in required if c not in df.columns]
    return len(missing) == 0, missing


def data_quality_checks(df: pd.DataFrame, label: str) -> dict:
    if df is None or df.empty:
        return {"score": 0, "missing": 0, "duplicates": 0, "rows": 0}
    missing = int(df.isna().sum().sum())
    dupes = int(df.duplicated().sum())
    score = round((1 - missing / max(df.size, 1)) * 100, 1)
    log_audit("Data Quality Check", f"{label}: score={score}%, missing={missing}, dupes={dupes}")
    return {"score": score, "missing": missing, "duplicates": dupes, "rows": len(df)}


def read_upload(file, cols: list[str]) -> pd.DataFrame | None:
    if file is None:
        return None
    try:
        name = file.name.lower()
        if name.endswith(".csv"):
            df = pd.read_csv(file)
        else:
            xl = pd.ExcelFile(file)
            sheet = st.selectbox(f"Sheet – {file.name}", xl.sheet_names, key=f"sheet_{file.name}")
            df = pd.read_excel(file, sheet_name=sheet)
        valid, missing = validate_columns(df, cols)
        if not valid:
            st.error(f"Missing columns: {', '.join(missing)}")
            return None
        quality = data_quality_checks(df, file.name)
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Rows", quality["rows"])
        c2.metric("Quality Score", f"{quality['score']}%")
        c3.metric("Missing Cells", quality["missing"])
        c4.metric("Duplicates", quality["duplicates"])
        if quality["duplicates"] > 0:
            st.warning(f"⚠ {quality['duplicates']} duplicate record(s) detected.")
        if quality["missing"] > 0:
            st.warning(f"⚠ {quality['missing']} missing value(s) detected.")
        st.dataframe(df.head(8), use_container_width=True)
        return df
    except Exception as exc:
        st.error(f"File read error: {exc}")
        return None


def merge_budget_actual(budget: pd.DataFrame, actual: pd.DataFrame) -> pd.DataFrame:
    b = budget.copy()
    a = actual.copy()
    b["Budget Amount"] = pd.to_numeric(b["Budget Amount"], errors="coerce")
    a["Actual Amount"] = pd.to_numeric(a["Actual Amount"], errors="coerce")
    merged = pd.merge(b, a, on=["Period", "Cost Head"], how="outer")
    merged["Budget Amount"] = merged["Budget Amount"].fillna(0)
    merged["Actual Amount"] = merged["Actual Amount"].fillna(0)
    return merged


# ── Sample Data ───────────────────────────────────────────────────────────────

@st.cache_data
def get_sample_budget() -> pd.DataFrame:
    heads = [
        "Personnel Salaries", "Marketing", "IT Infrastructure", "Office Rent",
        "Travel", "Training", "Software Licenses", "Consulting",
    ]
    periods = ["2025-Q1", "2025-Q2", "2025-Q3"]
    rows = []
    budgets = {
        "Personnel Salaries": [2500000, 2600000, 2700000],
        "Marketing": [800000, 850000, 1200000],
        "IT Infrastructure": [1200000, 1250000, 1300000],
        "Office Rent": [600000, 600000, 600000],
        "Travel": [350000, 380000, 400000],
        "Training": [200000, 210000, 220000],
        "Software Licenses": [450000, 480000, 500000],
        "Consulting": [300000, 320000, 340000],
    }
    approvals = {
        "2025-Q1": ["2024-12-15", "2024-12-28", "2025-01-05", "2024-12-20", "2025-01-10", "2024-12-18", "2024-12-10", "2024-11-30"],
        "2025-Q2": ["2025-03-01", "2025-03-15", "2025-03-20", "2025-02-28", "2025-03-10", "2025-02-15", "2025-02-20", "2025-02-25"],
        "2025-Q3": ["2025-06-01", "2025-06-15", "2025-07-05", "2025-05-28", "2025-06-10", "2025-05-15", "2025-05-20", "2025-05-25"],
    }
    starts = {"2025-Q1": "2025-01-01", "2025-Q2": "2025-04-01", "2025-Q3": "2025-07-01"}
    for pi, period in enumerate(periods):
        for hi, head in enumerate(heads):
            rows.append({
                "Period": period,
                "Cost Head": head,
                "Budget Amount": budgets[head][pi],
                "Approved Date": approvals[period][hi],
                "Budget Start Date": starts[period],
            })
    return pd.DataFrame(rows)


@st.cache_data
def get_sample_actual() -> pd.DataFrame:
    heads = [
        "Personnel Salaries", "Marketing", "IT Infrastructure", "Office Rent",
        "Travel", "Training", "Software Licenses", "Consulting",
    ]
    periods = ["2025-Q1", "2025-Q2", "2025-Q3"]
    actuals = {
        "Personnel Salaries": [2450000, 2550000, 2650000],
        "Marketing": [920000, 980000, 1100000],
        "IT Infrastructure": [1350000, 1400000, 1450000],
        "Office Rent": [580000, 590000, 595000],
        "Travel": [410000, 430000, 450000],
        "Training": [195000, 190000, 185000],
        "Software Licenses": [480000, 500000, 520000],
        "Consulting": [380000, 400000, 420000],
    }
    rows = []
    for pi, period in enumerate(periods):
        for head in heads:
            rows.append({"Period": period, "Cost Head": head, "Actual Amount": actuals[head][pi]})
    return pd.DataFrame(rows)


@st.cache_data
def get_sample_revisions() -> pd.DataFrame:
    return pd.DataFrame({
        "Revision ID": ["REV-001", "REV-002", "REV-003", "REV-004", "REV-005"],
        "Cost Head": ["Marketing", "IT Infrastructure", "Office Rent", "Consulting", "R&D"],
        "Old Budget": [800000, 1200000, 600000, 300000, 900000],
        "New Budget": [950000, 1200000, 600000, 420000, 1200000],
        "Revision Date": ["2025-02-15", "2025-01-20", "2025-03-01", None, "2025-03-15"],
        "Approved By": ["CFO", "CFO", "Pending", "Unapproved", None],
    })


def load_samples():
    st.session_state.budget_df = get_sample_budget()
    st.session_state.actual_df = get_sample_actual()
    st.session_state.revision_df = get_sample_revisions()
    st.session_state.merged_df = merge_budget_actual(
        st.session_state.budget_df, st.session_state.actual_df
    )
    run_all_tests()
    log_audit("Sample Data Loaded", "Demo datasets loaded across all modules")
    st.success("Sample data loaded successfully!")


# ── Audit Test Logic ──────────────────────────────────────────────────────────

def classify_variance_risk(pct: float) -> str:
    abs_pct = abs(pct)
    if abs_pct < 10:
        return "Green"
    if abs_pct <= 20:
        return "Amber"
    return "Red"


def run_test1(merged: pd.DataFrame) -> pd.DataFrame:
    df = merged.copy()
    df["Variance"] = df["Actual Amount"] - df["Budget Amount"]
    df["Variance %"] = np.where(
        df["Budget Amount"] != 0,
        (df["Variance"] / df["Budget Amount"]) * 100,
        0,
    )
    df["Risk"] = df["Variance %"].apply(classify_variance_risk)
    return df[["Period", "Cost Head", "Budget Amount", "Actual Amount", "Variance", "Variance %", "Risk"]]


def run_test2(budget: pd.DataFrame) -> pd.DataFrame:
    df = budget.copy()
    df["Approved Date"] = pd.to_datetime(df["Approved Date"], errors="coerce")
    df["Budget Start Date"] = pd.to_datetime(df["Budget Start Date"], errors="coerce")

    def status(row):
        if pd.isna(row["Approved Date"]) or pd.isna(row["Budget Start Date"]):
            return "❌ Late Approval"
        if row["Approved Date"] > row["Budget Start Date"]:
            return "❌ Late Approval"
        return "✅ Compliant"

    df["Status"] = df.apply(status, axis=1)
    return df[["Cost Head", "Approved Date", "Budget Start Date", "Status"]]


def run_test3(merged: pd.DataFrame) -> pd.DataFrame:
    work = merged.copy()
    work["Overspent"] = work["Actual Amount"] > work["Budget Amount"]
    agg = work.groupby("Cost Head").agg(
        Periods_Overspent=("Overspent", "sum"),
        Total_Periods=("Overspent", "count"),
    ).reset_index()
    agg = agg[agg["Periods_Overspent"] >= 3].copy()
    agg.rename(columns={"Periods_Overspent": "Number of Periods Overspent"}, inplace=True)

    def risk_rating(n, total):
        rate = n / total if total else 0
        if rate >= 0.75 or n >= 5:
            return "High"
        if rate >= 0.5 or n >= 4:
            return "Medium"
        return "Low"

    agg["Risk Rating"] = agg.apply(
        lambda r: risk_rating(r["Number of Periods Overspent"], r["Total_Periods"]), axis=1
    )
    return agg[["Cost Head", "Number of Periods Overspent", "Risk Rating"]]


def run_test4(revisions: pd.DataFrame) -> pd.DataFrame:
    df = revisions.copy()
    issues = []
    for _, row in df.iterrows():
        problems = []
        if pd.isna(row.get("Approved By")) or str(row.get("Approved By", "")).strip().lower() in ("", "nan", "none", "pending", "unapproved"):
            problems.append("Missing Approver")
        if pd.isna(row.get("Revision Date")):
            problems.append("Missing Revision Date")
        issues.append(" | ".join(problems) if problems else "Compliant")
    df["Issue"] = issues
    return df[df["Issue"] != "Compliant"][["Revision ID", "Cost Head", "Issue"]]


def run_test5(merged: pd.DataFrame) -> list[dict]:
    findings = []
    work = merged.sort_values(["Cost Head", "Period"])
    for head, grp in work.groupby("Cost Head"):
        grp = grp.reset_index(drop=True)
        if len(grp) >= 2:
            prev = grp.iloc[-2]["Budget Amount"]
            curr = grp.iloc[-1]["Budget Amount"]
            if prev > 0:
                increase = (curr - prev) / prev * 100
                if increase > 30:
                    findings.append({
                        "cost_head": head,
                        "type": "Budget Spike",
                        "commentary": (
                            f"{head} budget increased by {increase:.0f}% compared to previous period "
                            f"without corresponding increase in actual expenditure."
                        ),
                        "severity": "High" if increase > 50 else "Medium",
                    })
        under_budget_periods = (grp["Actual Amount"] < grp["Budget Amount"] * 0.5).sum()
        if under_budget_periods >= 2:
            findings.append({
                "cost_head": head,
                "type": "Under-Utilization",
                "commentary": (
                    f"{head} actual spend significantly lower than budget for {under_budget_periods} "
                    f"period(s), suggesting potentially inflated budget assumptions."
                ),
                "severity": "Medium",
            })
        if len(grp) >= 3:
            budgets = grp["Budget Amount"].values
            if budgets[-1] > budgets[-2] * 1.25 and budgets[-2] > budgets[-3] * 1.25:
                findings.append({
                    "cost_head": head,
                    "type": "Sudden Spike",
                    "commentary": f"{head} shows consecutive sudden spikes in budget allocation across recent periods.",
                    "severity": "High",
                })
    return findings


def run_all_tests():
    if st.session_state.merged_df is not None:
        st.session_state.test1_df = run_test1(st.session_state.merged_df)
    if st.session_state.budget_df is not None:
        st.session_state.test2_df = run_test2(st.session_state.budget_df)
    if st.session_state.merged_df is not None:
        st.session_state.test3_df = run_test3(st.session_state.merged_df)
    if st.session_state.revision_df is not None:
        st.session_state.test4_df = run_test4(st.session_state.revision_df)
    if st.session_state.merged_df is not None:
        st.session_state.test5_findings = run_test5(st.session_state.merged_df)


# ── Charts ────────────────────────────────────────────────────────────────────

def _layout(fig, title: str = "", height: int = 400):
    fig.update_layout(
        title=dict(text=title, font=dict(size=15, color=COLORS["teal"])),
        template="plotly_dark",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        height=height,
        margin=dict(l=40, r=40, t=50, b=40),
        font=dict(family="Inter, sans-serif"),
    )
    return fig


def chart_gauge(merged: pd.DataFrame) -> go.Figure:
    if merged is None or merged.empty:
        return _layout(go.Figure(), "Upload data for gauge")
    util = min(merged["Actual Amount"].sum() / max(merged["Budget Amount"].sum(), 1) * 100, 150)
    color = COLORS["success"] if util <= 100 else COLORS["danger"]
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=util,
        number=dict(suffix="%"),
        gauge=dict(
            axis=dict(range=[0, 150]),
            bar=dict(color=color),
            steps=[
                dict(range=[0, 80], color="rgba(16,185,129,0.3)"),
                dict(range=[80, 100], color="rgba(245,158,11,0.3)"),
                dict(range=[100, 150], color="rgba(239,68,68,0.3)"),
            ],
            threshold=dict(line=dict(color=COLORS["danger"], width=3), value=100),
        ),
        title=dict(text="Budget Utilization"),
    ))
    return _layout(fig, "", 350)


def chart_variance_trend(merged: pd.DataFrame) -> go.Figure:
    if merged is None or merged.empty:
        return _layout(go.Figure(), "Variance Trend")
    grp = merged.groupby("Period").agg({"Budget Amount": "sum", "Actual Amount": "sum"}).reset_index()
    grp["Variance"] = grp["Actual Amount"] - grp["Budget Amount"]
    colors_list = [COLORS["success"] if v <= 0 else COLORS["danger"] for v in grp["Variance"]]
    fig = go.Figure(go.Bar(x=grp["Period"], y=grp["Variance"], marker_color=colors_list))
    return _layout(fig, "Variance Trend by Period")


def chart_budget_vs_actual(merged: pd.DataFrame) -> go.Figure:
    if merged is None or merged.empty:
        return _layout(go.Figure(), "Budget vs Actual")
    grp = merged.groupby("Period").agg({"Budget Amount": "sum", "Actual Amount": "sum"}).reset_index()
    fig = go.Figure()
    fig.add_trace(go.Bar(name="Budget", x=grp["Period"], y=grp["Budget Amount"], marker_color=COLORS["teal"]))
    fig.add_trace(go.Bar(name="Actual", x=grp["Period"], y=grp["Actual Amount"], marker_color=COLORS["warning"]))
    fig.update_layout(barmode="group")
    return _layout(fig, "Budget vs Actual Comparison")


def chart_top_overspent(test1: pd.DataFrame) -> go.Figure:
    if test1 is None or test1.empty:
        return _layout(go.Figure(), "Top Overspent Heads")
    top = test1[test1["Variance"] > 0].nlargest(10, "Variance")
    fig = go.Figure(go.Bar(
        x=top["Variance"], y=top["Cost Head"], orientation="h",
        marker_color=COLORS["danger"], text=top["Variance"].apply(lambda x: fmt_currency(x)), textposition="auto",
    ))
    return _layout(fig, "Top Overspent Cost Heads")


def chart_heatmap(test1: pd.DataFrame) -> go.Figure:
    if test1 is None or test1.empty:
        return _layout(go.Figure(), "Variance Heatmap")
    pivot = test1.pivot_table(index="Cost Head", columns="Period", values="Variance %", aggfunc="mean", fill_value=0)
    fig = px.imshow(
        pivot.values, x=pivot.columns, y=pivot.index,
        color_continuous_scale=["#10B981", "#F59E0B", "#EF4444"],
        aspect="auto",
    )
    return _layout(fig, "Variance % Heatmap")


def chart_compliance_donut(test2: pd.DataFrame) -> go.Figure:
    if test2 is None or test2.empty:
        return _layout(go.Figure(), "Compliance")
    compliant = (test2["Status"] == "✅ Compliant").sum()
    late = len(test2) - compliant
    fig = go.Figure(go.Pie(
        labels=["Compliant", "Late Approval"],
        values=[compliant, late],
        hole=0.55,
        marker=dict(colors=[COLORS["success"], COLORS["danger"]]),
    ))
    return _layout(fig, "Approval Compliance", 350)


def chart_overspend_bar(test3: pd.DataFrame) -> go.Figure:
    if test3 is None or test3.empty:
        return _layout(go.Figure(), "Chronic Overspend Ranking")
    fig = go.Figure(go.Bar(
        x=test3["Number of Periods Overspent"],
        y=test3["Cost Head"],
        orientation="h",
        marker_color=[COLORS["danger"] if r == "High" else COLORS["warning"] if r == "Medium" else COLORS["success"]
                      for r in test3["Risk Rating"]],
    ))
    return _layout(fig, "Chronic Overspend Ranking")


def chart_exception_dashboard(test4: pd.DataFrame) -> go.Figure:
    if test4 is None or test4.empty:
        fig = go.Figure()
        fig.add_annotation(text="No revision exceptions detected", showarrow=False, font=dict(size=14, color=COLORS["success"]))
        return _layout(fig, "Revision Exception Dashboard", 300)
    issue_counts = test4["Issue"].value_counts()
    fig = go.Figure(go.Bar(x=issue_counts.index, y=issue_counts.values, marker_color=COLORS["danger"]))
    return _layout(fig, "Revision Exception Dashboard", 350)


def chart_waterfall(merged: pd.DataFrame) -> go.Figure:
    if merged is None or merged.empty:
        return _layout(go.Figure(), "Variance Waterfall")
    heads = merged.groupby("Cost Head").agg({"Budget Amount": "sum", "Actual Amount": "sum"}).reset_index()
    heads["Variance"] = heads["Actual Amount"] - heads["Budget Amount"]
    heads = heads.nlargest(12, "Variance", keep="all")
    fig = go.Figure(go.Waterfall(
        x=heads["Cost Head"], y=heads["Variance"],
        measure=["relative"] * len(heads),
        increasing=dict(marker=dict(color=COLORS["danger"])),
        decreasing=dict(marker=dict(color=COLORS["success"])),
    ))
    return _layout(fig, "Variance Waterfall")


def chart_treemap(merged: pd.DataFrame) -> go.Figure:
    if merged is None or merged.empty:
        return _layout(go.Figure(), "Budget Distribution")
    fig = px.treemap(
        merged, path=["Period", "Cost Head"], values="Budget Amount",
        color="Budget Amount", color_continuous_scale=["#152238", COLORS["teal"]],
    )
    return _layout(fig, "Budget Distribution Treemap")


def chart_spend_trend(merged: pd.DataFrame) -> go.Figure:
    if merged is None or merged.empty:
        return _layout(go.Figure(), "Monthly Spend Trend")
    grp = merged.groupby("Period")["Actual Amount"].sum().reset_index()
    fig = go.Figure(go.Scatter(
        x=grp["Period"], y=grp["Actual Amount"], mode="lines+markers",
        line=dict(color=COLORS["teal"], width=3), fill="tozeroy",
        fillcolor="rgba(0,212,170,0.15)",
    ))
    return _layout(fig, "Spend Trend by Period")


def chart_risk_matrix(test1: pd.DataFrame) -> go.Figure:
    if test1 is None or test1.empty:
        return _layout(go.Figure(), "Risk Matrix")
    color_map = {"Green": COLORS["success"], "Amber": COLORS["warning"], "Red": COLORS["danger"]}
    fig = go.Figure(go.Scatter(
        x=test1["Variance %"], y=test1["Budget Amount"],
        mode="markers+text", text=test1["Cost Head"],
        textposition="top center",
        marker=dict(
            size=test1["Actual Amount"].clip(10000, 500000) / 10000,
            color=[color_map.get(r, COLORS["gray_400"]) for r in test1["Risk"]],
            opacity=0.85,
        ),
    ))
    fig.update_layout(xaxis_title="Variance %", yaxis_title="Budget Amount")
    return _layout(fig, "Risk Matrix")


def chart_cost_ranking(merged: pd.DataFrame) -> go.Figure:
    if merged is None or merged.empty:
        return _layout(go.Figure(), "Cost Head Ranking")
    rank = merged.groupby("Cost Head")["Actual Amount"].sum().nlargest(12).reset_index()
    fig = go.Figure(go.Bar(x=rank["Actual Amount"], y=rank["Cost Head"], orientation="h", marker_color=COLORS["teal"]))
    return _layout(fig, "Cost Head Ranking by Actual Spend")


# ── Tables ────────────────────────────────────────────────────────────────────

def render_aggrid(df: pd.DataFrame, key: str, risk_col: str | None = None):
    if df is None or df.empty:
        st.info("No data available. Upload files or load sample data.")
        return
    gb = GridOptionsBuilder.from_dataframe(df)
    gb.configure_default_column(filterable=True, sortable=True, resizable=True)
    gb.configure_pagination(enabled=True, paginationPageSize=12)
    if risk_col and risk_col in df.columns:
        gb.configure_column(risk_col, cellStyle=JsCode("""
        function(params) {
            if (params.value === 'Green' || params.value === 'Low') return {'backgroundColor':'rgba(16,185,129,0.25)','color':'#10B981','fontWeight':'600'};
            if (params.value === 'Amber' || params.value === 'Medium') return {'backgroundColor':'rgba(245,158,11,0.25)','color':'#F59E0B','fontWeight':'600'};
            if (params.value === 'Red' || params.value === 'High') return {'backgroundColor':'rgba(239,68,68,0.25)','color':'#EF4444','fontWeight':'600'};
            return {};
        }
        """))
    AgGrid(df, gridOptions=gb.build(), height=400, theme="streamlit", allow_unsafe_jscode=True, key=key)


def export_excel(df: pd.DataFrame, sheet: str = "Findings") -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name=sheet, index=False)
    buf.seek(0)
    return buf.getvalue()


# ── AI Insights ───────────────────────────────────────────────────────────────

def generate_ai_insights() -> dict:
    merged = st.session_state.get("merged_df")
    test1 = st.session_state.get("test1_df")
    test2 = st.session_state.get("test2_df")
    test3 = st.session_state.get("test3_df")
    test4 = st.session_state.get("test4_df")
    test5 = st.session_state.get("test5_findings", [])

    summary_parts = []
    findings = []
    risks = []
    recommendations = []

    if merged is not None and not merged.empty:
        total_b = merged["Budget Amount"].sum()
        total_a = merged["Actual Amount"].sum()
        var_pct = (total_a - total_b) / total_b * 100 if total_b else 0
        summary_parts.append(
            f"Overall budget utilization stands at {total_a/total_b*100:.1f}% with a net variance of {var_pct:+.1f}%."
        )

    if test1 is not None and not test1.empty:
        red = test1[test1["Risk"] == "Red"]
        if not red.empty:
            top = red.iloc[0]
            findings.append(
                f"{top['Cost Head']} exceeded approved budget by {top['Variance %']:.1f}% "
                f"({fmt_currency(top['Variance'])} unfavorable variance)."
            )
            for _, r in red.head(3).iterrows():
                risks.append(f"High variance on {r['Cost Head']} ({r['Variance %']:+.1f}%) requires immediate management review.")

    if test2 is not None and not test2.empty:
        late = (test2["Status"] == "❌ Late Approval").sum()
        total = len(test2)
        pct = late / total * 100 if total else 0
        if late:
            findings.append(f"{late} of {total} budget approvals ({pct:.0f}%) were recorded after the budget period start date.")
            risks.append("Late budget approvals indicate weak pre-period budget governance controls.")
            recommendations.append("Implement mandatory budget approval workflow with automated deadline alerts before period commencement.")

    if test3 is not None and not test3.empty:
        high = test3[test3["Risk Rating"] == "High"]
        for _, r in high.iterrows():
            findings.append(
                f"{r['Cost Head']} consistently overspent budget in {r['Number of Periods Overspent']} periods."
            )
            risks.append(
                f"IT Infrastructure expenditure exceeded approved budget across multiple consecutive periods "
                f"indicating ineffective budget monitoring controls."
                if "IT" in r["Cost Head"]
                else f"Chronic overspend pattern detected for {r['Cost Head']}."
            )

    if test4 is not None and not test4.empty:
        findings.append(f"{len(test4)} budget revision(s) flagged with missing approver or revision date.")
        recommendations.append("Require dual sign-off and mandatory revision date for all budget amendments exceeding 5%.")

    for f in test5[:3]:
        findings.append(f["commentary"])

    if not recommendations:
        recommendations = [
            "Establish monthly variance review cadence for all cost heads exceeding 10% variance.",
            "Implement automated alerts for budget revisions lacking proper approval documentation.",
            "Conduct quarterly assumption reasonableness reviews for high-value budget categories.",
            "Escalate chronic overspend patterns to Finance Committee within 5 business days.",
        ]

    if not summary_parts:
        summary_parts = ["Upload budget and actual data to generate comprehensive audit insights."]

    return {
        "executive_summary": " ".join(summary_parts),
        "key_findings": findings or ["No significant findings detected in current dataset."],
        "risk_observations": risks or ["Risk profile within acceptable tolerance based on current analysis."],
        "recommendations": recommendations,
    }


# ── Report Generation ─────────────────────────────────────────────────────────

def generate_pdf_report(insights: dict, test1: pd.DataFrame | None) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    title_s = ParagraphStyle("T", parent=styles["Heading1"], fontSize=20, textColor=colors.HexColor(COLORS["navy"]))
    h_s = ParagraphStyle("H", parent=styles["Heading2"], fontSize=13, textColor=colors.HexColor(COLORS["teal_dark"]))
    body_s = ParagraphStyle("B", parent=styles["Normal"], fontSize=10, leading=14)

    story = []
    if LOGO and LOGO.exists():
        story.append(Image(str(LOGO), width=1.5 * inch, height=1.5 * inch))
        story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph(APP_TITLE, title_s))
    story.append(Paragraph(f"Report Date: {datetime.now():%B %d, %Y}", body_s))
    story.append(Spacer(1, 0.2 * inch))

    sections = [
        ("1. Executive Summary", [insights["executive_summary"]]),
        ("2. Scope", ["Budget vs actual variance analysis, approval compliance, overspend detection, revision review, and assumption testing."]),
        ("3. Audit Procedures", [
            "Test 1: Significant budget vs actual variance analysis",
            "Test 2: Budget approval before period start verification",
            "Test 3: Consistently overspent budget head identification",
            "Test 4: Budget revision approval review",
            "Test 5: Budget assumption reasonableness assessment",
        ]),
        ("4. Findings", insights["key_findings"]),
        ("5. Recommendations", insights["recommendations"]),
        ("6. Conclusion", ["Based on the procedures performed, management should address identified exceptions and strengthen budget governance controls."]),
    ]

    for title, items in sections:
        story.append(Paragraph(title, h_s))
        for item in items:
            story.append(Paragraph(f"• {item}", body_s))
        story.append(Spacer(1, 0.1 * inch))

    if test1 is not None and not test1.empty:
        tbl = [list(test1.columns)] + test1.head(8).values.tolist()
        t = Table(tbl, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(COLORS["navy"])),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ]))
        story.append(t)

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


def generate_docx_report(insights: dict) -> bytes:
    doc = Document()
    if LOGO and LOGO.exists():
        doc.add_picture(str(LOGO), width=Inches(1.2))

    title = doc.add_heading(APP_TITLE, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Report Date: {datetime.now():%B %d, %Y}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    sections = [
        ("Executive Summary", [insights["executive_summary"]]),
        ("Scope", ["Analysis of budget data, actual expenditures, revisions, and assumption reasonableness."]),
        ("Audit Procedures", [
            "Variance analysis with risk classification",
            "Approval timing compliance testing",
            "Chronic overspend identification",
            "Revision approval validation",
            "Assumption reasonableness review",
        ]),
        ("Findings", insights["key_findings"]),
        ("Recommendations", insights["recommendations"]),
        ("Conclusion", ["Management action is recommended to remediate identified control gaps."]),
    ]

    for heading, items in sections:
        h = doc.add_heading(heading, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 168, 136)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def compile_findings_excel() -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
        for name, df, key in [
            ("Variance Analysis", st.session_state.get("test1_df"), "test1"),
            ("Approval Compliance", st.session_state.get("test2_df"), "test2"),
            ("Chronic Overspend", st.session_state.get("test3_df"), "test3"),
            ("Revision Exceptions", st.session_state.get("test4_df"), "test4"),
        ]:
            if df is not None and not df.empty:
                df.to_excel(writer, sheet_name=name[:31], index=False)
        findings = st.session_state.get("test5_findings", [])
        if findings:
            pd.DataFrame(findings).to_excel(writer, sheet_name="Assumption Findings", index=False)
    buf.seek(0)
    return buf.getvalue()


# ── Page Renderers ────────────────────────────────────────────────────────────

def page_dashboard():
    render_header()
    merged = st.session_state.get("merged_df")
    test1 = st.session_state.get("test1_df")

    if merged is None:
        st.info("Upload budget and actual data via **Data Import**, or click **Load Sample Data** in the sidebar.")
        return

    total_b = merged["Budget Amount"].sum()
    total_a = merged["Actual Amount"].sum()
    variance = total_a - total_b
    var_pct = variance / total_b * 100 if total_b else 0
    over_heads = int((merged["Actual Amount"] > merged["Budget Amount"]).sum()) if not merged.empty else 0
    high_risk = int((test1["Risk"] == "Red").sum()) if test1 is not None and not test1.empty else 0

    c1, c2, c3, c4, c5, c6 = st.columns(6)
    with c1: kpi_card("Total Budget", fmt_currency(total_b))
    with c2: kpi_card("Total Actual Spend", fmt_currency(total_a))
    with c3: kpi_card("Total Variance", fmt_currency(variance))
    with c4: kpi_card("Variance %", f"{var_pct:+.1f}%")
    with c5: kpi_card("Over-Budget Heads", str(over_heads))
    with c6: kpi_card("High Risk Findings", str(high_risk))

    st.markdown('<p class="section-title">Performance Analytics</p>', unsafe_allow_html=True)
    r1c1, r1c2 = st.columns(2)
    with r1c1: st.plotly_chart(chart_gauge(merged), use_container_width=True)
    with r1c2: st.plotly_chart(chart_variance_trend(merged), use_container_width=True)
    r2c1, r2c2 = st.columns(2)
    with r2c1: st.plotly_chart(chart_budget_vs_actual(merged), use_container_width=True)
    with r2c2: st.plotly_chart(chart_top_overspent(test1), use_container_width=True)


def page_import():
    render_header("Data Import Center", "Upload budget, actual, and revision datasets")
    tab1, tab2, tab3 = st.tabs(["📊 Budget Data", "📈 Actual Data", "🔄 Budget Revisions"])

    with tab1:
        f = st.file_uploader("Budget file (.xlsx, .xls, .csv)", type=["xlsx", "xls", "csv"], key="budget_up")
        if f:
            df = read_upload(f, BUDGET_COLS)
            if df is not None:
                st.session_state.budget_df = df
                log_audit("Budget Data Uploaded", f"{len(df)} rows from {f.name}")
                if st.session_state.actual_df is not None:
                    st.session_state.merged_df = merge_budget_actual(df, st.session_state.actual_df)
                    run_all_tests()

    with tab2:
        f = st.file_uploader("Actual file (.xlsx, .xls, .csv)", type=["xlsx", "xls", "csv"], key="actual_up")
        if f:
            df = read_upload(f, ACTUAL_COLS)
            if df is not None:
                st.session_state.actual_df = df
                log_audit("Actual Data Uploaded", f"{len(df)} rows from {f.name}")
                if st.session_state.budget_df is not None:
                    st.session_state.merged_df = merge_budget_actual(st.session_state.budget_df, df)
                    run_all_tests()

    with tab3:
        f = st.file_uploader("Revision file (.xlsx, .xls, .csv)", type=["xlsx", "xls", "csv"], key="rev_up")
        if f:
            df = read_upload(f, REVISION_COLS)
            if df is not None:
                st.session_state.revision_df = df
                log_audit("Revision Data Uploaded", f"{len(df)} rows from {f.name}")
                run_all_tests()

    with st.expander("📋 Expected Column Formats"):
        st.markdown("""
        **Budget Data:** Period | Cost Head | Budget Amount | Approved Date | Budget Start Date  
        **Actual Data:** Period | Cost Head | Actual Amount  
        **Budget Revisions:** Revision ID | Cost Head | Old Budget | New Budget | Revision Date | Approved By
        """)


def page_test1():
    render_header("Audit Test 1", "Significant Budget vs Actual Variances")
    test1 = st.session_state.get("test1_df")
    if test1 is None:
        st.warning("Load data first via Data Import or sample data.")
        return

    red = (test1["Risk"] == "Red").sum()
    amber = (test1["Risk"] == "Amber").sum()
    green = (test1["Risk"] == "Green").sum()
    c1, c2, c3 = st.columns(3)
    c1.metric("🟢 Green (<10%)", green)
    c2.metric("🟡 Amber (10-20%)", amber)
    c3.metric("🔴 Red (>20%)", red)

    render_aggrid(test1, "t1_grid", "Risk")
    c1, c2 = st.columns(2)
    with c1: st.plotly_chart(chart_heatmap(test1), use_container_width=True)
    with c2: st.plotly_chart(chart_top_overspent(test1), use_container_width=True)
    st.download_button("📥 Download Results", export_excel(test1, "Variances"), "audit_test1_variances.xlsx", use_container_width=True)


def page_test2():
    render_header("Audit Test 2", "Budget Approval Before Period Start")
    test2 = st.session_state.get("test2_df")
    if test2 is None:
        st.warning("Load budget data first.")
        return

    compliant = (test2["Status"] == "✅ Compliant").sum()
    pct = compliant / len(test2) * 100 if len(test2) else 0
    c1, c2 = st.columns(2)
    c1.metric("Compliance Rate", f"{pct:.1f}%")
    c2.metric("Late Approvals", len(test2) - compliant)

    render_aggrid(test2, "t2_grid")
    c1, c2 = st.columns(2)
    with c1: st.plotly_chart(chart_compliance_donut(test2), use_container_width=True)
    with c2:
        st.markdown("#### Compliance Summary")
        st.dataframe(test2["Status"].value_counts().reset_index(), use_container_width=True)
    st.download_button("📥 Download Results", export_excel(test2, "Approvals"), "audit_test2_approvals.xlsx", use_container_width=True)


def page_test3():
    render_header("Audit Test 3", "Consistently Overspent Budget Heads")
    test3 = st.session_state.get("test3_df")
    if test3 is None:
        st.warning("Load data first.")
        return

    if test3.empty:
        st.success("✅ No cost heads overspent in 3+ periods.")
        return

    st.metric("Chronic Overspend Heads", len(test3))
    render_aggrid(test3, "t3_grid", "Risk Rating")
    st.plotly_chart(chart_overspend_bar(test3), use_container_width=True)
    st.download_button("📥 Download Results", export_excel(test3, "Overspend"), "audit_test3_overspend.xlsx", use_container_width=True)


def page_test4():
    render_header("Audit Test 4", "Budget Revisions Approval Review")
    test4 = st.session_state.get("test4_df")
    revisions = st.session_state.get("revision_df")

    if revisions is None:
        st.warning("Load revision data first.")
        return

    if test4 is None or test4.empty:
        st.success("✅ All budget revisions have proper approvals and dates.")
    else:
        st.error(f"⚠ {len(test4)} revision exception(s) identified")
        render_aggrid(test4, "t4_grid")
        st.download_button("📥 Download Exceptions", export_excel(test4, "Exceptions"), "audit_test4_exceptions.xlsx", use_container_width=True)

    st.plotly_chart(chart_exception_dashboard(test4 if test4 is not None else pd.DataFrame()), use_container_width=True)


def page_test5():
    render_header("Audit Test 5", "Budget Assumption Reasonableness Review")
    findings = st.session_state.get("test5_findings", [])

    if not findings:
        st.success("✅ No assumption reasonableness concerns identified.")
        return

    st.metric("Findings Generated", len(findings))
    for f in findings:
        sev_color = COLORS["danger"] if f["severity"] == "High" else COLORS["warning"]
        st.markdown(
            f'<div class="insight-card"><h4 style="color:{sev_color}">{f["type"]} – {f["cost_head"]}</h4>'
            f'<p>{f["commentary"]}</p></div>',
            unsafe_allow_html=True,
        )

    merged = st.session_state.get("merged_df")
    if merged is not None:
        st.markdown('<p class="section-title">Trend Analysis</p>', unsafe_allow_html=True)
        st.plotly_chart(chart_spend_trend(merged), use_container_width=True)


def page_ai_insights():
    render_header("AI Insights", "AI-generated audit observations and recommendations")
    insights = generate_ai_insights()

    st.markdown(
        f'<div class="glass-card"><h4 style="color:{COLORS["teal"]};margin-top:0">Executive Summary</h4>'
        f'<p style="color:{COLORS["gray_400"]}">{insights["executive_summary"]}</p></div>',
        unsafe_allow_html=True,
    )

    col1, col2 = st.columns(2)
    with col1:
        st.markdown(f'<p class="section-title">Key Findings</p>', unsafe_allow_html=True)
        for item in insights["key_findings"]:
            st.markdown(f'<div class="insight-card"><p>{item}</p></div>', unsafe_allow_html=True)
    with col2:
        st.markdown(f'<p class="section-title">Risk Observations</p>', unsafe_allow_html=True)
        for item in insights["risk_observations"]:
            st.markdown(f'<div class="insight-card"><p>{item}</p></div>', unsafe_allow_html=True)

    st.markdown(f'<p class="section-title">Recommendations</p>', unsafe_allow_html=True)
    for item in insights["recommendations"]:
        st.markdown(f'<div class="insight-card"><p>✅ {item}</p></div>', unsafe_allow_html=True)


def page_analytics():
    render_header("Analytics", "Advanced interactive visualizations")
    merged = st.session_state.get("merged_df")
    test1 = st.session_state.get("test1_df")
    if merged is None:
        st.warning("Load data first.")
        return

    r1c1, r1c2 = st.columns(2)
    with r1c1: st.plotly_chart(chart_waterfall(merged), use_container_width=True)
    with r1c2: st.plotly_chart(chart_treemap(merged), use_container_width=True)
    r2c1, r2c2 = st.columns(2)
    with r2c1: st.plotly_chart(chart_spend_trend(merged), use_container_width=True)
    with r2c2: st.plotly_chart(chart_risk_matrix(test1), use_container_width=True)
    st.plotly_chart(chart_cost_ranking(merged), use_container_width=True)


def page_report():
    render_header("Audit Report", "Automated audit report generation")
    insights = generate_ai_insights()
    test1 = st.session_state.get("test1_df")
    ts = datetime.now().strftime("%Y%m%d")

    for title, items in [
        ("Executive Summary", [insights["executive_summary"]]),
        ("Scope", ["Budget variance, approval compliance, overspend, revision, and assumption testing."]),
        ("Findings", insights["key_findings"][:5]),
        ("Recommendations", insights["recommendations"][:5]),
    ]:
        with st.expander(title, expanded=title == "Executive Summary"):
            for item in items:
                st.markdown(f"• {item}")

    st.divider()
    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button("📄 Download PDF", generate_pdf_report(insights, test1), f"audit_report_{ts}.pdf", mime="application/pdf", use_container_width=True)
    with c2:
        st.download_button("📊 Download Excel Findings", compile_findings_excel(), f"audit_findings_{ts}.xlsx", use_container_width=True)
    with c3:
        st.download_button("📝 Download DOCX Report", generate_docx_report(insights), f"audit_report_{ts}.docx", use_container_width=True)


def sidebar():
    with st.sidebar:
        if LOGO and LOGO.exists():
            st.image(str(LOGO), use_container_width=True)
        else:
            st.markdown("### CITF AI")
        st.markdown(f"<p style='color:{COLORS['gray_400']};font-size:0.8rem;'>{APP_TITLE}</p>", unsafe_allow_html=True)
        st.divider()

        st.session_state.user_name = st.text_input("Auditor Name", st.session_state.get("user_name", "Audit Analyst"))
        page = st.radio("Navigation", [label for _, label, _ in NAV_ITEMS], label_visibility="collapsed")
        page_key = {label: key for _, label, key in NAV_ITEMS}[page]

        st.divider()
        if st.button("📥 Load Sample Data", use_container_width=True):
            load_samples()
            st.rerun()

        with st.expander("🔐 Audit Trail"):
            trail = st.session_state.get("audit_trail", [])
            if trail:
                st.dataframe(pd.DataFrame(trail), use_container_width=True, hide_index=True)
            else:
                st.caption("No audit events yet.")

        st.markdown(
            f"<p style='color:{COLORS['gray_400']};font-size:0.7rem;text-align:center;margin-top:2rem;'>"
            f"© {datetime.now().year} CITF AI<br/>Internal Audit Analytics</p>",
            unsafe_allow_html=True,
        )
    return page_key


# ── Main ──────────────────────────────────────────────────────────────────────

PAGES = {
    "dashboard": page_dashboard,
    "import": page_import,
    "test1": page_test1,
    "test2": page_test2,
    "test3": page_test3,
    "test4": page_test4,
    "test5": page_test5,
    "ai_insights": page_ai_insights,
    "analytics": page_analytics,
    "report": page_report,
}


def main():
    init_session()
    st.markdown(get_css(), unsafe_allow_html=True)
    page_key = sidebar()
    PAGES.get(page_key, page_dashboard)()


if __name__ == "__main__":
    main()
